import streamlit as st
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
import tempfile
import os
from dotenv import load_dotenv
from pathlib import Path
import contextlib
from typing import Dict, Any, List
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.utils import formatdate
import subprocess
import logging
from pathlib import Path
import olefile 
from email.mime.application import MIMEApplication
from datetime import datetime
from config.settings import get_openai_api_key

# ログ設定
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# メール送信関連の定数を環境変数から取得
SMTP_SERVER = os.getenv('SMTP_SERVER', 'mail.alt-g.jp')
SMTP_PORT = int(os.getenv('SMTP_PORT', '587'))
SMTP_TIMEOUT = int(os.getenv('SMTP_TIMEOUT', '30'))
SMTP_TO_ADDRESS = os.getenv('SMTP_TO_ADDRESS', 'kanrieigyo@lberc-g.jp')
SMTP_DOMAIN = os.getenv('SMTP_DOMAIN', 'alt-g.jp')

def get_setting():
    """設定を取得する"""
    settings = {}
    
    # Cloud Run環境かどうかを確認
    is_cloud_run = os.getenv('CLOUD_RUN_ENV') == 'true'
    
    # 環境変数から設定を取得
    settings['OPENAI_API_KEY'] = os.getenv('OPENAI_API_KEY')
    settings['LABEL_EMAIL'] = os.getenv('LABEL_EMAIL')
    settings['LABEL_PASSWORD'] = os.getenv('LABEL_PASSWORD')
    
    # Cloud Run環境でない場合はStreamlit Secretsも確認
    if not is_cloud_run and hasattr(st, 'secrets'):
        try:
            if not settings['OPENAI_API_KEY']:
                settings['OPENAI_API_KEY'] = st.secrets.get('OPENAI_API_KEY')
            if not settings['LABEL_EMAIL']:
                settings['LABEL_EMAIL'] = st.secrets.get('LABEL_EMAIL')
            if not settings['LABEL_PASSWORD']:
                settings['LABEL_PASSWORD'] = st.secrets.get('LABEL_PASSWORD')
        except Exception as e:
            logger.warning(f"Streamlit Secretsからの読み込み中にエラー: {str(e)}")
    
    return settings

# 契約書種類ごとのプロンプト定義
CONTRACT_TYPES = {
    "基本契約書": {
        "prompt": """あなたは企業の法務担当者です。
    企業の契約書類を分析し、法務分析をしてださい。
    #指示
    ## 甲または乙の顧客に対する直接契約、商流飛ばしに関する記述がないか
       - ある場合は、賠償金規定について確認し、また甲乙双方が平等になっているかを確認してください
    ## 甲または乙の従業員に対する勧誘、引き抜き行為、または直接の雇用形態や準委任契約締結に関する記述がないか
       - ある場合は、賠償金規定について確認し、また甲乙双方が平等になっているかを確認してください
    ## 損害賠償について、甲乙が平等な契約になっているか
    ## 合意管轄裁判所が"東京地方裁判所"になっているか。
       - 東京地方裁判所でない場合は、東京地方裁判所への変更、あるいは専属的を削除し、合意管轄裁判所に変更
    ## 契約解除に関する規定で、合意的理由がなく甲が一方的に契約を解除できる規定がないか
    
    #Example
    ##input
    第１２条（解約）
    甲は、本契約期間中であって、１か月前の予告をすることにより、本契約を解約することができる

    ##output
    変更例：
    甲および乙は、本契約期間中であっても、相手方に対して少なくとも１か月前に書面による予告をすることにより、正当な理由に基づいて本契約を解約することができる。
    正当な理由には以下が含まれるが、これに限定されない
    1. 契約違反
    2. 経済的困難
    3.その他双方が合意する理由 

    ##input
    （損害賠償）
    第１９条 乙が自己の責に帰するべき事由により甲に損害を与えた場合、本契約又は個別契約もしくは注文書による注文の解除、解約の有無に関わらずただちに当該損害を賠償すべき責めを負うものとする。

    ##output
    第１９条 "甲または乙"が自己の責に帰するべき事由により"相手方"に損害を与えた場合、本契約又は個別契約もしくは注文書による注文の解除、解約の有無に関わらずただちに当該損害を賠償すべき責めを負うものとする。

    #input
    甲及び乙は、本契約及び個別契約に関する訴訟については、"大阪地方裁判所"を第1審の専属的合意管轄裁判所とする。
    #output
    甲及び乙は、本契約及び個別契約に関する訴訟については、"東京地方裁判所"を第1審の専属的合意管轄裁判所とする。
    もしくは、甲及び乙は、本契約及び個別契約に関する訴訟については、大阪地方裁判所を第1審の"合意管轄裁判所"とする。

    #input
    第12条 （直接交渉等の制限）
    1. 乙は、本契約期間中及び、契約終了後 2 年間は、甲の書面による 事前の承諾を得ることなく、
    甲から業務委託された案件における甲の顧客（甲に対する業務委託者 等をいい、以下「甲顧客」
    という。）との間で、直接又は第三者を通じて、当該案件に関するコンピュータ ソフトウェアに
    係る業務又はこれに付随する業務その他当該案件を含む甲の顧客におけるプロ ジェクトに関
    わる契約交渉を行ってはならない。
    2. 乙が前項に違反して甲顧客との間で直接又は第三者を通じて契約を締 結した場合、乙は、当該
    契約に基づく甲の受注金額の３倍の額を、違約金として、甲に対して支 払わなければならない。
    なお、本項の定めは、これを超える損害が甲に発生した場合における乙の 損害賠償責任を免除
    するものではない。 

    ##output
    第12条 （直接交渉等の制限）
    1. 甲又は乙は、本契約期間中及び、契約終了後 2 年間は、相手方の書面による 事前の承諾を得ることなく、
    甲又は乙から業務委託された案件における相手方の顧客（相手方に対する業務委託者 等をいい、以下「相手方顧客」
    という。）との間で、直接又は第三者を通じて、当該案件に関するコンピュータ ソフトウェアに
    係る業務又はこれに付随する業務その他当該案件を含む相手方顧客におけるプロ ジェクトに関
    わる契約交渉を行ってはならない。
    2. 甲又は乙が前項に違反して相手方顧客との間で直接又は第三者を通じて契約を締 結した場合、甲又は乙は、当該
    契約に基づく相手方の受注金額の３倍の額を、違約金として、相手方に対して支 払わなければならない。
    なお、本項の定めは、これを超える損害が甲または乙に発生した場合におけ相手方の 損害賠償責任を免除
    するものではない。 

    ##input
    ３．乙は、本契約または個別契約の契約期間中及び契約終了後、甲の従業員、���注先な
    ど（以下、併せて、「従業員等」という。）を勧誘して委託業務と同種又は類似する
    業務を依頼する等の引抜行為をしてはならない。これに違反した場合、乙は、甲に
    対し、直ちに、既に甲から受領した全ての委託料の返還に加えて、違約罰として従
    業員等１人に金１００万円、その他別途甲に生じた損害を賠償するものとする。

    ##output
    ３．甲又は乙は、本契約または個別契約の契約期間中及び契約終了後、相手方の従業員、外注先な
    ど（以下、併せて、「従業員等」という。）を勧誘して委託業務と同種又は類似する
    業務を依頼する等の引抜行為をしてはならない。これに違反した場合、甲又は乙は、相手方に
    対し、直ちに、既に相手方から受領した全ての委託料の返還に加えて、違約罰として従
    業員等１人に金１００万円、その他別途相手方に生じた損害を賠償するものとする。"""
    },
    "機密保持契約書": {
        "prompt": """あなたは契約書分析の専門家です。機密保持契約書を分析し、以下の項目について詳細に明してください。
        分析項目：
    ## 損害賠償について、甲乙が平等な契約になっているか
    ## 合意管轄裁判所が"東京地方裁判所"になっているか。
       - 東京地方裁判所でない場合は、東京地方裁判所への変更、あるいは専属的を削除し、合意管轄裁判所に変更
    
    #Example
    ##input
    （損害賠償）
    第１９条 乙が自己の責に帰するべき事由により甲に損害を与えた場合、本契約又は個別契約もしくは注文書による注文の解除、解約の有無に関わらずただちに当該損害を賠償すべき責めを負うものとする。

    ##output
    第１９条 "甲または乙"が自己の責に帰するべき事由により"相手方"に損害を与えた場合、本契約又は個別契約もしくは注文書による注文の解除、解約の有無に関わらずただちに当該損害を賠償すべき責めを負うものとする。

        #input
    甲及び乙は、本契約及び個別契約に関する訴訟については、"大阪地方裁判所"を第1審の専属的合意管轄裁判所とする。
    #output
    甲及び乙は、本契約及び個別契約に関する訴訟については、"東京地方裁判所"を第1審の専属的合意管轄裁判所とする。
    もしくは、甲及び乙は、本契約及び個別契約に関する訴訟については、大阪地方裁判所を第1審の"合意管轄裁判所"とする。
    """
    },
}

FILENAME_ANALYSIS_PROMPT = """あなたは契約書分析の専門家です。
ファイル名から契約書の種類を判別してください。

以下の選択肢から最も適切な契約書の種類を1つ選んでください：
- 基本契約書
- 機密保持契約書
- その他

判断基準：
- 基本契約書:基本、基本契約書、業務委託基本契約書、SES基本契約書
- 機密保持契約書:NDA、情報漏洩防止
- その他：それ以外の反社会的勢力、出向契約、誓約書など

回答は契約書の種類のみを出力してください。理由や説明は不要です。

#Example
##input
【新】業務請負基本契約書(2019.09改定)
##output
基本契約書

##input
機密保持に関する覚書
##output
機密保持契約書

##input
反社会勢力両面印刷_令和版
##output
その他

ファイル名：{filename}"""

def detect_contract_type_from_filename(filename: str) -> str:
    """LLMを使用してファイル名から契約書の種類を判別"""
    model = ChatOpenAI(
        model="gpt-4o-mini",
        temperature=0
    )
    
    chain = (
        ChatPromptTemplate.from_template(FILENAME_ANALYSIS_PROMPT)
        | model
        | StrOutputParser()
    )
    
    try:
        contract_type = chain.invoke({"filename": filename})
        return contract_type.strip()
    except Exception as e:
        st.error(f"契約書種類の判別中にエラーが発生しました: {str(e)}")
        return "その他"

def extract_text_from_doc(file_path: str) -> str:
    """
    olefileを使用して.docファイルからテキストを抽出する改善版
    """
    try:
        if not olefile.isOleFile(file_path):
            logger.warning(f"{file_path} is not a valid OLE file")
            return None

        ole = olefile.OleFileIO(file_path)
        try:
            # 利用可能なストリームを確認
            streams = ole.listdir()
            logger.info(f"Available streams: {streams}")

            # WordDocumentストリームの処理
            if ole.exists('WordDocument'):
                word_stream = ole.openstream('WordDocument').read()
                
                # テキストデータの位置を特定するための異なるエンコーディングを試行
                encodings = ['utf-16le', 'utf-8', 'shift-jis', 'euc-jp']
                extracted_text = None
                
                for encoding in encodings:
                    try:
                        # バイナリデータからテキストを抽出する試み
                        text = word_stream.decode(encoding, errors='ignore')
                        # 基本的なクリーニング
                        cleaned_text = ''.join(char for char in text if char.isprintable() or char in '\n\r\t')
                        # 最小限の有効性チェック
                        if len(cleaned_text.strip()) > 0 and any(c.isalnum() for c in cleaned_text):
                            extracted_text = cleaned_text
                            logger.info(f"Successfully extracted text using {encoding} encoding")
                            break
                    except Exception as e:
                        logger.debug(f"Failed to decode with {encoding}: {str(e)}")
                        continue

                if extracted_text:
                    # 追加のクリーニング処理
                    # 複数の空白行を1つに置換
                    cleaned_text = '\n'.join(line for line in extracted_text.splitlines() if line.strip())
                    return cleaned_text

            # Tablesストリームの処理（補助的なテキスト抽出）
            if ole.exists('Table'):
                table_stream = ole.openstream('Table').read()
                try:
                    table_text = table_stream.decode('utf-16le', errors='ignore')
                    if table_text.strip():
                        return table_text.strip()
                except Exception as e:
                    logger.debug(f"Failed to extract text from Table stream: {str(e)}")

            logger.warning("No text could be extracted from the document streams")
            return None

        finally:
            ole.close()

    except Exception as e:
        logger.error(f"Error in extract_text_from_doc: {str(e)}")
        return None

def process_uploaded_file(uploaded_file):
    """アップロードされたファイルを処理"""
    try:
        with safe_temp_file(suffix=Path(uploaded_file.name).suffix) as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            file_ext = Path(uploaded_file.name).suffix.lower()
            
            file_type = {
                '.pdf': 'pdf',
                '.doc': 'doc',
                '.docx': 'docx'
            }.get(file_ext)
            
            if not file_type:
                st.error("サポートされていないファイル形式です")
                return None
                
            extracted_text = extract_document_text(tmp_file.name, file_type)
            return extracted_text
                
    except Exception as e:
        logger.error(f"Error processing file {uploaded_file.name}: {str(e)}", exc_info=True)
        st.error(f"ファイル処理中にエラーが発生しました: {str(e)}")
        return None

def extract_document_text(file_path: str, file_type: str) -> str:
    """ドキュメントからテキストを抽出する改善版"""
    try:
        if file_type == 'pdf':
            loader = PyPDFLoader(file_path)
            pages = loader.load()
            return "\n".join(page.page_content for page in pages)
            
        elif file_type == 'doc':
            # まずolefileでの抽出を試みる
            text = extract_text_from_doc(file_path)
            if text and len(text.strip()) > 0:
                logger.info("Successfully extracted text using olefile")
                return text
                
            # olefileでの抽出が失敗した場合、追加の方法を試行
            logger.info("Attempting alternative extraction methods")
            
            # docx2txtを試行
            try:
                loader = Docx2txtLoader(file_path)
                documents = loader.load()
                extracted_text = "\n".join(doc.page_content for doc in documents)
                if extracted_text.strip():
                    logger.info("Successfully extracted text using docx2txt")
                    return extracted_text
            except Exception as e:
                logger.debug(f"docx2txt extraction failed: {str(e)}")

            # python-docxを試行
            try:
                from docx import Document
                doc = Document(file_path)
                extracted_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                if extracted_text.strip():
                    logger.info("Successfully extracted text using python-docx")
                    return extracted_text
            except Exception as e:
                logger.debug(f"python-docx extraction failed: {str(e)}")

            logger.error("All text extraction methods failed")
            return None
            
        elif file_type == 'docx':
            try:
                # まずpython-docxを試す
                from docx import Document
                doc = Document(file_path)
                extracted_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                if extracted_text.strip():
                    logger.info("Successfully extracted text using python-docx")
                    return extracted_text
            except Exception as e:
                logger.debug(f"python-docx extraction failed: {str(e)}")
                
            # python-docxが失敗した場合はdocx2txtを試す
            try:
                loader = Docx2txtLoader(file_path)
                documents = loader.load()
                extracted_text = "\n".join(doc.page_content for doc in documents)
                if extracted_text.strip():
                    logger.info("Successfully extracted text using docx2txt")
                    return extracted_text
            except Exception as e:
                logger.debug(f"docx2txt extraction failed: {str(e)}")
                
            logger.error("All text extraction methods failed")
            return None
            
        else:
            logger.error(f"Unsupported file type: {file_type}")
            return None
            
    except Exception as e:
        logger.error(f"Text extraction error: {str(e)}")
        st.error(f"テキスト抽出中にエラーが発生しました: {str(e)}")
        return None

def analyze_document(text: str, contract_type: str, company_name: str) -> str:
    """契約書の分析を実行"""
    if contract_type not in CONTRACT_TYPES:
        return "この種類の契約書は分析対象外です。"
        
    prompt = CONTRACT_TYPES[contract_type]["prompt"]
    
    template = f"""
    以下の契約書を分析してくだ��い：

    契約当事者: {company_name}
    
    契約書テキスト:
    {{text}}
    """

    chain = (
        ChatPromptTemplate.from_messages([
            ("system", prompt),
            ("human", template)
        ])
        | ChatOpenAI(model="gpt-4o-mini", temperature=0, streaming=True)
        | StrOutputParser()
    )
    
    try:
        return chain.invoke({"text": text})
    except Exception as e:
        return f"分析中にエラーが発生しました: {str(e)}"

@contextlib.contextmanager
def safe_temp_file(suffix=None):
    """一時ファイルを安全に扱うためのコンテキストマネージャー"""
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    try:
        yield temp_file
    finally:
        temp_file.close()
        if os.path.exists(temp_file.name):
            os.unlink(temp_file.name)

@st.dialog('リーガルチェック依頼')
def confirm_send_dialog(contract_company, source_company, duration, uploaded_files):
    """メール送信確認ダイアログ"""
    st.write("以下の内容で承認依頼を送信します。内容を確認してください。")
    st.write(f"上位会社名: {contract_company}")
    st.write(f"契約元会社: {source_company}")
    st.write(f"期限: {duration}")
    st.write(f"添付ファイル数: {len(uploaded_files)}件")
    
    # ボタンを横に並べるためのカラムを作成
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button('送信する', type="primary", use_container_width=True):
            st.session_state.start_analysis = True
            st.session_state.show_dialog = False
            st.rerun()
    
    with col2:
        if st.button('キャンセル', use_container_width=True):
            st.session_state.show_dialog = False
            st.rerun()

def legal_check_page():
    """契約書分析ページのメイン関数"""
    # 設定を取得
    settings = get_setting()
    
    # 必要な設定が揃っているか確認
    missing_settings = []
    if not settings.get('OPENAI_API_KEY'):
        missing_settings.append("OPENAI_API_KEY")
    if not settings.get('LABEL_EMAIL'):
        missing_settings.append("LABEL_EMAIL")
    if not settings.get('LABEL_PASSWORD'):
        missing_settings.append("LABEL_PASSWORD")
    
    if missing_settings:
        st.error(f"以下の設定が見つかりません：\n" + 
                "\n".join([f"- {setting}" for setting in missing_settings]) + 
                "\n\n環境変数または.streamlit/secrets.tomlで設定してください。")
        st.info("""
        設定方法：
        1. 環境変数として設定：
           ```
           set OPENAI_API_KEY=your-api-key
           set LABEL_EMAIL=your-email
           set LABEL_PASSWORD=your-password
           ```
        
        2. または.streamlit/secrets.tomlに記述：
           ```
           OPENAI_API_KEY = "your-api-key"
           LABEL_EMAIL = "your-email"
           LABEL_PASSWORD = "your-password"
           ```
        """)
        return
    
    # OpenAI APIキーの確認
    if not settings.get('OPENAI_API_KEY'):
        st.warning("OpenAI APIキーが設定されていません。環境変数OPENAI_API_KEYを設定してください。")
        return
    
    # グローバル変数として定義
    global LABEL_EMAIL, LABEL_PASSWORD
    LABEL_EMAIL = settings.get('LABEL_EMAIL')
    LABEL_PASSWORD = settings.get('LABEL_PASSWORD')

    if not LABEL_EMAIL or not LABEL_PASSWORD:
        logger.error("メール設定が見つかりません")
        st.error("メール設定が正しく構成されていません。システム管理者に連絡してください。")
        return

    # セッションステートの初期化
    if 'show_dialog' not in st.session_state:
        st.session_state.show_dialog = False
    
    st.title("📄リーガルチェック")
    
    # 入力フォーム
    col1, col2 = st.columns(2)
    with col1:
        contract_company = st.text_input("上位会社名", placeholder="上位会社名")
        contract_type = st.selectbox(
            "契約理由", 
            ["新規契約", "契約書変更", "機密保持のみ契約", "その他"]
        )
        
    with col2:
        source_company = st.selectbox(
            "契約元会社名",
             ["ACR","LBQ","LBA","LBG","VLB","LBJ"]
        )
        duration = st.selectbox(
            "期限", 
            ["5営業日", "3営業日", "至急", "期限なし"]
        )
    
    # ファイルアップロード
    st.header("2. 契約書ファイルのアップロード")
    uploaded_files = st.file_uploader(
        "契約書をアップロード（複数選択可）",
        type=["pdf", "docx", "doc"],
        accept_multiple_files=True,
        help="PDFまたはWord形式のファイル（1ファイルあたり最大50MB）"
    )

    # 承認依頼ボタン
    if st.button("承認依頼", use_container_width=True, type="primary"):
        if not contract_company:
            st.error("契約法人名を入力してください")
            return
        
        if not uploaded_files:
            st.error("契約書ファイルをアップロードしてください")
            return
            
        # ダイアログ表示フラグを設定
        st.session_state.show_dialog = True
        st.rerun()
    
    # ダイアログの表示制御
    if st.session_state.show_dialog:
        confirm_send_dialog(contract_company, source_company, duration, uploaded_files)
    
    # 分析とメール送信の自動実行
    if st.session_state.get('start_analysis', False):
        # フラグをリセット
        st.session_state.start_analysis = False
        
        # 分析処理
        results = []
        processable_files = []
        skipped_files = []
        
        # プログレスバーのコンテナを作成
        progress_container = st.empty()
        
        # まず、すべてのファイルの種類を判別
        with st.spinner("契約書の種類を判別中..."):
            for file in uploaded_files:
                file_contract_type = detect_contract_type_from_filename(file.name)
                if file_contract_type == "その他":
                    skipped_files.append(file.name)
                else:
                    processable_files.append((file, file_contract_type))
            
            if not processable_files:
                st.error("分析可能な契約書（基本契約書または機密保持契約書）が見つかりませんでした。")
                return
            
            # 分析���能なファイルの処理
            for file, file_contract_type in processable_files:
                try:
                    extracted_text = process_uploaded_file(file)
                    if extracted_text:
                        # 分析実行
                        analysis = analyze_document(
                            extracted_text,
                            file_contract_type,
                            contract_company
                        )
                        
                        results.append({
                            "filename": file.name,
                            "contract_type": file_contract_type,
                            "analysis": analysis
                        })
                
                except Exception as e:
                    logger.error(f"Error processing file {file.name}: {str(e)}", exc_info=True)
                    results.append({
                        "filename": file.name,
                        "contract_type": "エラー",
                        "analysis": f"処理中にエラーが発生しました: {str(e)}"
                    })
            
            # 分析完了後、自動的にメール送信
            try:
                if send_legal_check_email(
                    company_name=contract_company,
                    source_company=source_company,
                    contract_type=contract_type,
                    duration=duration,
                    results=results,
                    uploaded_files=uploaded_files
                ):
                    st.success("メールを送信しました")
                else:
                    st.error("メール送信に失敗しました")
            except Exception as e:
                logger.error(f"メール送信中にエラーが発生: {str(e)}")
                st.error(f"メール送信中にエラーが発生しました: {str(e)}")

def create_analysis_text(
    company_name: str,
    source_company: str,
    contract_type: str,
    duration: str,
    results: List[Dict]
) -> str:
    """分析結果のテキストを作成"""
    text = f"""リーガルチェック分析結果

============= 
上位名　：{company_name}
契約元　：{source_company}
契約事由：{contract_type}
期　限　：{duration}
分析日時：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
=============

"""
    if results:
        for result in results:
            text += f"\n【ファイル名】{result['filename']}\n"
            text += f"【契約書種類】{result['contract_type']}\n"
            text += f"【分析結果】\n{result['analysis']}\n"
            text += "=" * 50 + "\n"
    
    return text

def send_legal_check_email(
    company_name: str,
    source_company: str,
    contract_type: str,
    duration: str,
    results: List[Dict],
    uploaded_files: List
) -> bool:
    """リーガルチェック依頼のメール送信（添付ファイル付き）"""
    try:
        # Secretsからメール設定を取得
        settings = get_setting()
        smtp_user = settings.get('LABEL_EMAIL')
        smtp_password = settings.get('LABEL_PASSWORD')
        
        if not smtp_user or not smtp_password:
            logger.error("メール送信設定が見つかりません")
            st.error("メール送信の設定が正しくありません。システム管理者に連絡してください。")
            return False

        # ユーザー情報の取得
        user_name = st.session_state.get('name', '未設定')
        
        # メールアドレスの正規化（ユーザー名のみを使用）
        login_user = smtp_user.lower().strip()
        if '@' in login_user:
            login_user = login_user.split('@')[0]
        
        from_addr = f"{login_user}@{SMTP_DOMAIN}"  # 完全なメールアドレス
            
        # 設定情報のログ出力
        logger.info("=== メール送信設定情報 ===")
        logger.info(f"SMTP Server: {SMTP_SERVER}")
        logger.info(f"SMTP Port: {SMTP_PORT}")
        logger.info(f"Login User: {login_user}")
        logger.info(f"From Address: {from_addr}")
        logger.info(f"To Address: {SMTP_TO_ADDRESS}")
        logger.info(f"User Name: {user_name}")
        
        # メールメッセージの作成
        msg = MIMEMultipart()
        msg['Subject'] = f"【リーガルチェック依頼】【{duration}】{company_name}"
        msg['From'] = from_addr  # 完全なメールアドレスを使用
        msg['To'] = SMTP_TO_ADDRESS
        msg['Date'] = formatdate(localtime=True)
        
        # メール本文
        body = f"""自動送信メール

リーガルチェック依頼がきております。
承認担当者は期限までにご確認お願いいたします。

============= 
上位名　：{company_name}
契約元　：{source_company}
契約事由：{contract_type}
期　限　：{duration}
担当者　：{user_name}
=============
"""
        msg.attach(MIMEText(body, 'plain', 'utf-8'))

        # 分析結果の添付
        if results:
            analysis_text = create_analysis_text(
                company_name,
                source_company,
                contract_type,
                duration,
                results
            )
            analysis_attachment = MIMEText(analysis_text, 'plain', 'utf-8')
            analysis_attachment.add_header(
                'Content-Disposition',
                'attachment',
                filename='分析結果.txt'
            )
            msg.attach(analysis_attachment)

        # 元の契約書ファ��ルを添付
        for uploaded_file in uploaded_files:
            file_attachment = MIMEApplication(
                uploaded_file.getvalue(),
                _subtype=Path(uploaded_file.name).suffix[1:]  # 拡張子から.を除去
            )
            file_attachment.add_header(
                'Content-Disposition',
                'attachment',
                filename=uploaded_file.name
            )
            msg.attach(file_attachment)

        # メールヘッダー情報のログ
        logger.info("=== メールヘッダー情報 ===")
        for header, value in msg.items():
            logger.info(f"{header}: {value}")

        # SMTPサーバーへの接続とメール送信
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT, timeout=SMTP_TIMEOUT) as smtp_server:
            smtp_server.set_debuglevel(2)
            
            # EHLO/HELO
            logger.info("=== SMTP接続開始 ===")
            try:
                ehlo_response = smtp_server.ehlo(SMTP_DOMAIN)
                logger.info(f"EHLO Response: {ehlo_response}")
            except Exception as e:
                logger.error(f"EHLO失敗: {str(e)}")
                raise
            
            # STARTTLS
            if smtp_server.has_extn('STARTTLS'):
                logger.info("STARTTLSを開始")
                smtp_server.starttls()
                ehlo_response = smtp_server.ehlo(SMTP_DOMAIN)
                logger.info(f"STARTTLS後のEHLO Response: {ehlo_response}")
            
            # ログイン（ユーザー名のみを使用）
            try:
                logger.info(f"SMTPログイン試行 - ユーザー: {login_user}")
                smtp_server.login(login_user, smtp_password)
                logger.info("SMTPログイン成功")
            except smtplib.SMTPAuthenticationError as auth_error:
                logger.error(f"SMTP認証エラー: {str(auth_error)}")
                logger.error(f"エラーコード: {auth_error.smtp_code}")
                logger.error(f"エラーメッセージ: {auth_error.smtp_error.decode('utf-8') if hasattr(auth_error, 'smtp_error') else ''}")
                raise
            
            # メール送信
            try:
                logger.info("=== メール送信試行 ===")
                smtp_server.sendmail(
                    from_addr=from_addr,
                    to_addrs=[SMTP_TO_ADDRESS],
                    msg=msg.as_string()
                )
                logger.info("メール送信成功")
                return True
                
            except smtplib.SMTPSenderRefused as e:
                logger.error("=== 送信者アドレス拒否エラー ===")
                logger.error(f"エラーコード: {e.smtp_code}")
                logger.error(f"エラーメッセージ: {e.smtp_error.decode('utf-8') if hasattr(e, 'smtp_error') else str(e)}")
                logger.error(f"送信者アドレス: {e.sender}")
                st.error(f"送信者アドレスが拒否されました: {e.smtp_error.decode('utf-8') if hasattr(e, 'smtp_error') else str(e)}")
                return False
                
    except Exception as e:
        logger.error("=== メール送信処理でエラー発生 ===")
        logger.error(f"エラータイプ: {type(e).__name__}")
        logger.error(f"エラーメッセージ: {str(e)}")
        logger.error("詳細なスタックトレース:", exc_info=True)
        if hasattr(e, 'smtp_code'):
            logger.error(f"SMTPエラーコード: {e.smtp_code}")
        if hasattr(e, 'smtp_error'):
            logger.error(f"SMTPエラーメッセージ: {e.smtp_error.decode('utf-8') if isinstance(e.smtp_error, bytes) else e.smtp_error}")
        st.error("メール送信処理中に予期せぬエラーが発生しました。")
        return False

if __name__ == "__main__":
    legal_check_page()